"""Vacation order (atostogų įsakymas) generation for UAB People Link."""
import io
import logging
import re
from datetime import date, timedelta

import httpx
import openpyxl
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.config import settings
from app.sharepoint import resolve_token_and_drive

logger = logging.getLogger(__name__)

MONTH_NOM = {
    1: "SAUSIS", 2: "VASARIS", 3: "KOVAS", 4: "BALANDIS",
    5: "GEGUŽĖ", 6: "BIRŽELIS", 7: "LIEPA", 8: "RUGPJŪTIS",
    9: "RUGSĖJIS", 10: "SPALIS", 11: "LAPKRITIS", 12: "GRUODIS",
}
MONTH_GEN = {
    1: "sausio", 2: "vasario", 3: "kovo", 4: "balandžio",
    5: "gegužės", 6: "birželio", 7: "liepos", 8: "rugpjūčio",
    9: "rugsėjo", 10: "spalio", 11: "lapkričio", 12: "gruodžio",
}

HOLIDAYS_2026 = {
    date(2026, 1, 1), date(2026, 2, 16), date(2026, 3, 11),
    date(2026, 5, 1), date(2026, 6, 24), date(2026, 7, 6),
    date(2026, 8, 15), date(2026, 11, 1), date(2026, 11, 2),
    date(2026, 12, 25), date(2026, 12, 26),
}

EXCEL_PATH = "Atostogų prašymai.xlsx"
EXCEL_DRIVE = "Praymai"
ORDERS_FOLDER = "Atostogos/Įsakymai_Atostogos_2026"
ORDERS_DRIVE = "sakymai"
LEAVE_ORDER = ["kasmetines", "mama", "teva", "stazas", "neapmokamos"]


def first_working_day(year: int, month: int) -> date:
    d = date(year, month, 1)
    while d.weekday() >= 5 or d in HOLIDAYS_2026:
        d += timedelta(days=1)
    return d


def to_dative(name: str) -> str:
    """Best-effort Lithuanian nominative → dative for common name endings."""
    parts = name.strip().split()
    result = []
    for p in parts:
        converted = p
        for suffix, replacement in [
            ("ienė", "ienei"), ("aitė", "aitei"), ("ytė", "ytei"),
            ("ė", "ei"), ("aitis", "aitiui"), ("ytis", "yčiui"),
            ("as", "ui"), ("is", "iui"), ("us", "ui"), ("ys", "iui"),
            ("a", "ai"),
        ]:
            if p.lower().endswith(suffix):
                converted = p[:-len(suffix)] + replacement
                break
        result.append(converted)
    return " ".join(result)


def days_text(n: int) -> str:
    if n == 1:
        return "1 darbo dieną"
    elif 2 <= n <= 9:
        return f"{n} darbo dienas"
    return f"{n} darbo dienų"


def get_leave_type(comment: str) -> str:
    c = comment.strip().lower()
    if not c:
        return "kasmetines"
    if "mamadienis" in c:
        return "mama"
    if "tėvadienis" in c or "tevadienis" in c:
        return "teva"
    if "stažą" in c or "stazą" in c or "stažu" in c:
        return "stazas"
    if "neapmokamos" in c:
        return "neapmokamos"
    return "kasmetines"


_GRAPH = "https://graph.microsoft.com/v1.0"


async def _get_drive(url_name: str) -> tuple[str, str]:
    """Return (token, drive_base) for the library whose webUrl ends with /{url_name}."""
    from app.sharepoint import _get_token, _cache_get, _cache_set

    cache_key = f"vac:drive:{url_name}"
    token = await _get_token(
        settings.sharepoint_tenant_id,
        settings.sharepoint_client_id,
        settings.sharepoint_client_secret,
    )
    cached = _cache_get(cache_key)
    if cached:
        return token, cached

    auth = {"Authorization": f"Bearer {token}"}
    async with httpx.AsyncClient(timeout=30) as client:
        site_resp = await client.get(
            f"{_GRAPH}/sites/{settings.sharepoint_site_hostname}:/{settings.sharepoint_site_path}",
            headers=auth,
        )
        site_resp.raise_for_status()
        site_id = site_resp.json()["id"]

        drives_resp = await client.get(f"{_GRAPH}/sites/{site_id}/drives", headers=auth)
        drives_resp.raise_for_status()
        drives = drives_resp.json()["value"]

        drive = next(
            (d for d in drives if d.get("webUrl", "").rstrip("/").endswith(f"/{url_name}")),
            None,
        )
        if not drive:
            raise RuntimeError(f"Library '{url_name}' not found on SharePoint")
        drive_base = f"{_GRAPH}/drives/{drive['id']}"

    _cache_set(cache_key, drive_base, 86400)
    return token, drive_base


async def download_excel() -> bytes:
    token, drive_base = await _get_drive(EXCEL_DRIVE)
    async with httpx.AsyncClient(timeout=60, follow_redirects=True) as client:
        resp = await client.get(
            f"{drive_base}/root:/{EXCEL_PATH}:/content",
            headers={"Authorization": f"Bearer {token}"},
        )
        resp.raise_for_status()
        return resp.content


async def get_next_order_number() -> int:
    token, drive_base = await _get_drive(ORDERS_DRIVE)
    async with httpx.AsyncClient(timeout=30) as client:
        resp = await client.get(
            f"{drive_base}/root:/{ORDERS_FOLDER}:/children",
            headers={"Authorization": f"Bearer {token}"},
            params={"$select": "name", "$top": "500"},
        )
        if not resp.is_success:
            logger.warning(f"Could not list orders folder: {resp.status_code}")
            return 1
        names = [item["name"] for item in resp.json().get("value", [])]
    max_num = 0
    for name in names:
        m = re.search(r"P-(\d+)", name)
        if m:
            max_num = max(max_num, int(m.group(1)))
    return max_num + 1


def parse_excel(content: bytes, month_num: int) -> list[dict]:
    wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
    ws = wb.active

    # Find row index of each month header
    month_rows: dict[int, int] = {}  # row_idx -> month_num
    for row_idx, row in enumerate(ws.iter_rows(), 1):
        for cell in row:
            val = str(cell.value or "").strip().upper()
            for mn, mname in MONTH_NOM.items():
                if val == mname and mn not in month_rows.values():
                    month_rows[row_idx] = mn

    target_row = next((r for r, m in month_rows.items() if m == month_num), None)
    if target_row is None:
        return []

    subsequent = sorted(r for r in month_rows if r > target_row)
    end_row = subsequent[0] if subsequent else ws.max_row + 1

    entries = []
    header_skipped = False

    for row_idx in range(target_row + 1, end_row):
        row_vals = [c.value for c in ws[row_idx]]
        if not any(v for v in row_vals):
            continue

        # Skip the column header row (contains "Vardas", "Nr." etc.)
        if not header_skipped:
            row_str = " ".join(str(v or "") for v in row_vals)
            if any(kw in row_str for kw in ("Vardas", "Nr.", "Eil")):
                header_skipped = True
                continue

        name = str(row_vals[1] or "").strip() if len(row_vals) > 1 else ""
        if not name:
            continue

        terminas = row_vals[3] if len(row_vals) > 3 else ""
        days_val = row_vals[4] if len(row_vals) > 4 else 0
        comment = str(row_vals[6] or "").strip() if len(row_vals) > 6 else ""

        if "IKI" in comment.upper():
            continue

        try:
            days_n = int(float(str(days_val or 0)))
        except (ValueError, TypeError):
            days_n = 0

        entries.append({
            "name": name,
            "name_dative": to_dative(name),
            "date_str": _format_terminas(terminas),
            "days": days_n,
            "leave_type": get_leave_type(comment),
        })

    return entries


def _format_terminas(val) -> str:
    if not val:
        return ""
    if isinstance(val, str):
        return val.strip()
    if hasattr(val, "month"):  # openpyxl datetime
        return f"{val.year} m. {MONTH_GEN[val.month]} {val.day} d."
    return str(val).strip()


def process_entries(entries: list[dict]) -> list[dict]:
    """Merge same person + same leave type; sort by type order."""
    grouped: dict[tuple, list] = {}
    for e in entries:
        key = (e["name"], e["leave_type"])
        grouped.setdefault(key, []).append(e)

    result = []
    for (name, leave_type), group in grouped.items():
        if len(group) == 1:
            result.append(group[0])
        else:
            merged = dict(group[0])
            merged["periods"] = group
            result.append(merged)

    order_map = {t: i for i, t in enumerate(LEAVE_ORDER)}
    return sorted(result, key=lambda e: order_map.get(e["leave_type"], 99))


def build_entry_text(entry: dict, month_num: int) -> str:
    lt = entry["leave_type"]
    mon = MONTH_GEN[month_num]
    periods = entry.get("periods", [entry])
    name_d = entry["name_dative"]

    if lt == "kasmetines":
        parts = [
            f"{days_text(p['days'])} kasmetinių atostogų {'2026 ' if i == 0 else ''}{p['date_str']}"
            for i, p in enumerate(periods)
        ]
        return f"{name_d} {' ir '.join(parts)}, atostoginius sumokant kartu su {mon} mėnesio atlyginimu."

    if lt in ("mama", "teva"):
        parts = [
            f"1 papildomą poilsio dieną {'2026 m. ' if i == 0 else ''}{p['date_str']}"
            for i, p in enumerate(periods)
        ]
        return f"{name_d} {' ir '.join(parts)}, nes augina vaiką iki dvylikos metų, atostoginius sumokant kartu su {mon} mėnesio atlyginimu."

    if lt == "stazas":
        dk = "(vadovaujantis Lietuvos Respublikos darbo kodekso (Žin., 2002, Nr. 64-2569; 2002, Nr. IX-926) 233 straipsniu)"
        parts = [
            f"{days_text(p['days'])} kasmetinių atostogų {dk} {'2026 ' if i == 0 else ''}{p['date_str']}"
            for i, p in enumerate(periods)
        ]
        return f"{name_d} {' ir '.join(parts)}, atostoginius sumokant kartu su {mon} mėnesio atlyginimu."

    if lt == "neapmokamos":
        parts = [
            f"{days_text(p['days'])} neapmokamų atostogų {'2026 ' if i == 0 else ''}{p['date_str']}"
            for i, p in enumerate(periods)
        ]
        return f"{name_d} {' ir '.join(parts)}."

    return ""


def generate_docx(entries: list[dict], month_num: int, order_num: int, order_date: date) -> bytes:
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    def add_para(text: str, bold: bool = False, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        run.bold = bold
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    add_para('UŽDAROSIOS AKCINĖS BENDROVĖS „PEOPLE LINK“', bold=True)
    add_para("DIREKTORĖ", bold=True)
    add_para("ĮSAKYMAS", bold=True)
    add_para("DĖL KASMETINIŲ, NEAPMOKAMŲ ATOSTOGŲ IR PAPILDOMŲ POILSIO DIENŲ SUTEIKIMO", bold=True)
    add_para("")
    add_para(f"2026 m. {MONTH_GEN[order_date.month]} {order_date.day} d.  Nr. P-{order_num}")
    add_para("Vilnius")
    add_para("")
    add_para("Nusprendžiau suteikti:", align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para("")

    for i, entry in enumerate(entries, 1):
        text = build_entry_text(entry, month_num)
        if not text:
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f"{i}. {text}")
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    add_para("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run1 = p.add_run("Direktorė")
    run1.font.name = "Times New Roman"
    run1.font.size = Pt(12)
    run2 = p.add_run("                                        Jurgita Lemešiūtė")
    run2.font.name = "Times New Roman"
    run2.font.size = Pt(12)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
